import PySimpleGUI as sg
from docx import Document
from docx.shared import Pt
import pandas as pd
from random import randint
from docx.shared import Pt
import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class TelaPython:
    def __init__(self):
        #layout
        layout = [
            [sg.Text('Nome da Empresa', size=(17, 0)), sg.Input(key='nomeempresa')],
            [sg.Text('Código de Barras', size=(17, 0)), sg.Input(key='codigobarras')],
            [sg.Text('CNPJ', size=(17, 0)), sg.Input(key='cnpj')],
            [sg.Text('Data de Validade', size=(17, 0)), sg.Input(key='datavalidade')],
            [sg.Text('Competência', size=(17, 0)), sg.Input(key='compt')],
            [sg.Text('Valor Recolhido', size=(17, 0)), sg.Input(key='valorrecolhido')],
            [sg.Text('Data de Pagto.', size=(17, 0)), sg.Input(key='datapag')],
            [sg.Text('Hora de Pgto.', size=(17, 0)), sg.Input(key='horapag')],
            [sg.Button('Enviar Dados')]

        ]
        #janela
        janela = sg.Window('Dados do Usuário').layout(layout)
        #Extrair os dados da tela
        self.button, self.values = janela.Read()





    def Iniciar(self):
        nomeempresa = self.values['nomeempresa']
        codigobarras = self.values['codigobarras']
        cnpj = self.values['cnpj']
        datavalidade = self.values['datavalidade']
        compt = self.values['compt']
        valorrecolhido = self.values['valorrecolhido']
        datapag = self.values['datapag']
        horapag = self.values['horapag']

        document = Document('Comprovante de pagamento - modelo(2).docx')

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)
        paragraph5 = document.add_paragraph('                                           Nome:')
        paragraph5.add_run(f'      {nomeempresa}').bold = True
        paragraph5.paragraph_format.space_before = Pt(0)
        paragraph5.paragraph_format.space_after = Pt(4)
        paragraph6 = document.add_paragraph()
        paragraph6.add_run('                                        Agência:     ')
        paragraph6.add_run(f'{randint(1000, 9999)}').bold = True
        paragraph6.add_run('                            Conta:    ')
        paragraph6.add_run(f'{randint(1000, 9999)}-{randint(0, 9)}').bold = True
        paragraph7 = document.add_paragraph()
        paragraph7.add_run('')
        paragraph8 = document.add_paragraph()
        paragraph8.add_run(f'                     Código de Barras:    {codigobarras}').bold = True
        paragraph8.paragraph_format.space_after = Pt(0)
        paragraph8.paragraph_format.space_after = Pt(4)
        paragraph9 = document.add_paragraph()
        paragraph9.add_run('                                         CNPJ:     ')
        paragraph9.add_run(f'{cnpj}').bold = True
        paragraph9.paragraph_format.space_after = Pt(4)
        paragraph10 = document.add_paragraph()
        paragraph10.add_run('                   Código de convênio:     ')
        paragraph10.add_run('0155').bold = True
        paragraph10.paragraph_format.space_after = Pt(4)
        paragraph11 = document.add_paragraph()
        paragraph11.add_run('                       Data de validade:     ')
        paragraph11.add_run(f'{datavalidade}').bold = True
        paragraph11.paragraph_format.space_after = Pt(4)
        paragraph12 = document.add_paragraph()
        paragraph12.add_run('                             Competência:     ')
        paragraph12.add_run(f'{compt}').bold = True
        paragraph12.paragraph_format.space_after = Pt(4)
        paragraph13 = document.add_paragraph()
        paragraph13.add_run('                          Valor recolhido:     ')
        paragraph13.add_run(f'R$ {valorrecolhido}').bold = True
        paragraph14 = document.add_paragraph()
        paragraph14.add_run(f' Pagamento efetuado em {datapag} às {horapag}h   via CEI, CTRL  1607080014567').bold = True
        paragraph15 = document.add_paragraph()
        paragraph15.paragraph_format.space_after = Pt(4)
        paragraph15.add_run(f'Autenticação:').bold = True
        paragraph15.paragraph_format.space_after = Pt(0)
        paragraph16 = document.add_paragraph()
        sentence = paragraph16.add_run(f'JSET7E3DBLQI7DO7E5L8P3D43XB4424FO0036583B14059440EE5D0B4E4BP3SP200000000020004320')
        sentence.font.size = Pt(8.5)

        document.save('C:\\Users\\user\\Documents\\Desktop\\Trabalho\\Comprovantes\\' + f'Comprovante de Pagamento - FGTS - {nomeempresa}.docx')


tela = TelaPython()
tela.Iniciar()